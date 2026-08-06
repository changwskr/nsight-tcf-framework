# -*- coding: utf-8 -*-
from docx import Document
from pathlib import Path

base = Path(__file__).resolve().parent
# find docx
docs = list(base.glob("*.docx"))
if not docs:
    raise SystemExit("no docx")
path = docs[0]
print("reading", path.name)
doc = Document(str(path))
out = base / "_extract.txt"
lines = []
lines.append(f"file={path.name}")
lines.append(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    t = (p.text or "").rstrip()
    if not t:
        continue
    lines.append(f"{i}|{style}|{t}")
lines.append("\n===== TABLES =====")
for ti, table in enumerate(doc.tables):
    lines.append(f"\n--- TABLE {ti} ({len(table.rows)}x{len(table.columns)}) ---")
    for ri, row in enumerate(table.rows):
        cells = [(c.text or "").replace("\n", " / ").strip() for c in row.cells]
        lines.append(f"{ri}| " + " || ".join(cells))
out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "bytes", out.stat().st_size)
